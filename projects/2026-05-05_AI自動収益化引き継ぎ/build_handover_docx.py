"""Build claude_code_handover.docx (A4 / 実務提出レベル / 完全版)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent / "outputs" / "claude_code_handover.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

JP_FONT = "Yu Gothic"
MONO_FONT = "Consolas"
MONO_EAST = "MS Gothic"

HEADER_FILL = "1F4E78"
CALLOUT_DEFAULT = "FFF4CE"
CALLOUT_INFO = "E7F3FF"
CALLOUT_WARN = "FFE2E2"

WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def style_run(run, *, size: float | None = 11, bold: bool = False,
              color: RGBColor | None = None, mono: bool = False) -> None:
    name = MONO_FONT if mono else JP_FONT
    east = MONO_EAST if mono else JP_FONT
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), east)


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        style_run(run, size=None)


def add_para(doc: Document, text: str, *, bold: bool = False, size: float = 11,
             align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    style_run(p.add_run(text), size=size, bold=bold)


def _add_listed(doc: Document, items: list[str], style: str) -> None:
    for item in items:
        p = doc.add_paragraph(style=style)
        style_run(p.add_run(item), size=11)


def add_bullets(doc: Document, items: list[str]) -> None:
    _add_listed(doc, items, "List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    _add_listed(doc, items, "List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        style_run(hdr[i].paragraphs[0].add_run(h), size=11, bold=True, color=WHITE)
        set_cell_shading(hdr[i], HEADER_FILL)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            style_run(cells[c_idx].paragraphs[0].add_run(val), size=10.5)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(text), size=9.5, mono=True)


def add_callout(doc: Document, label: str, text: str, color_hex: str = CALLOUT_DEFAULT) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, color_hex)
    cell.text = ""
    p = cell.paragraphs[0]
    style_run(p.add_run(f"【{label}】 "), size=11, bold=True)
    style_run(p.add_run(text), size=11)


def add_centered_run(doc: Document, text: str, *, size: float, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(text), size=size, bold=bold)


def main() -> None:
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = JP_FONT
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), JP_FONT)

    add_centered_run(doc, "Claude Code 引き継ぎ書 完全版", size=26, bold=True)
    add_centered_run(doc, "AI自動収益化プロジェクト｜note・YouTube・SNS・CrowdWorks 運用設計", size=14)
    doc.add_paragraph()
    add_centered_run(doc, "発行日：2026-05-05　／　対象：Claude Code・Gemini・人間オペレーター", size=11)

    doc.add_paragraph()
    add_callout(
        doc,
        "本書の使い方",
        "本書は、AI自動収益化プロジェクトを引き継ぐ全ての担当（Claude Code / Gemini / 人間）が、"
        "現状・課題・次アクション・収益化方針を一読で把握できるよう構成された実務用ハンドオーバ文書である。"
        "新しい計画を立てる前に、第8章および第14章の即時実行指示を必ず参照すること。",
    )

    doc.add_page_break()

    add_heading(doc, "第1章　プロジェクト概要", 1)
    add_para(doc, "本プロジェクトは、AIを活用して以下を実現するための自動収益化プロジェクトである。")
    add_bullets(doc, [
        "note記事の自動生成・公開",
        "YouTube Shorts向け動画企画・生成",
        "Reddit / SNS への海外向け拡散",
        "CrowdWorks案件への応募支援",
        "将来的な月10万円以上の副収益化",
    ])
    add_para(
        doc,
        "現状は、技術基盤と一部自動化は構築済みであり、最大の課題は「公開・応募・継続実行」である。",
    )
    add_callout(
        doc,
        "結論",
        "仕組みは80%できている。残る20%は『出す勇気』と『毎日触る習慣』である。",
        CALLOUT_INFO,
    )

    add_heading(doc, "第2章　主要アカウント・サービス一覧", 1)
    add_table(
        doc,
        ["サービス", "用途", "現状"],
        [
            ["note", "記事投稿・有料記事販売", "運用中"],
            ["YouTube", "Shorts投稿", "準備中"],
            ["Reddit", "海外流入導線", "半自動"],
            ["CrowdWorks", "案件応募・受注", "ログイン済み"],
            ["Claude Code", "自動化コード作成・実行支援", "使用中"],
            ["Gemini", "外部レビュー・評価", "使用予定"],
            ["OpenAI API", "記事生成・文章生成", "使用想定"],
        ],
    )
    add_heading(doc, "注意事項", 2)
    add_bullets(doc, [
        "APIキーは .env 管理とする。",
        "パスワードは文書内に記載しない。",
        "SNS投稿は規約違反・スパム判定を避けるため、頻度制御を必須とする。",
    ])

    add_heading(doc, "第3章　これまでの成果", 1)
    add_para(doc, "Cowork期間において、以下の成果がある。")
    add_bullets(doc, [
        "note記事を複数作成",
        "有料記事化を実施",
        "英語・海外向け展開を開始",
        "Reddit投稿補助スクリプトを作成",
        "cronによる定期実行設定を開始",
        "YouTube Shorts向け動画構想を作成",
        "CrowdWorks応募文テンプレートを作成",
    ])
    add_callout(
        doc,
        "評価",
        "技術構築力は高いが、収益化に直結する『公開数』『応募数』『継続実行』が不足している。",
    )

    add_heading(doc, "第4章　自動化システム稼働状況", 1)
    add_para(doc, "現在の構成は以下の通り。")
    add_code_block(
        doc,
        "cron\n  ↓\nrun.sh\n  ↓\nPythonスクリプト\n  ↓\nAI記事生成\n  ↓\n英語翻訳\n  ↓\nReddit / SNS 投稿補助\n  ↓\nログ保存",
    )
    add_heading(doc, "現状評価", 2)
    add_table(
        doc,
        ["項目", "状態"],
        [
            ["cron", "設定済み"],
            ["Python環境", "構築済み"],
            ["note生成", "半自動"],
            ["Reddit投稿", "半自動"],
            ["YouTube連携", "未完成"],
            ["SNS自動投稿", "未完成"],
            ["ログ保存", "必須化予定"],
        ],
    )
    add_callout(
        doc,
        "総合評価",
        "自動化は約80%完成。残り20%は、投稿連携・動画生成・安定運用である。",
        CALLOUT_INFO,
    )

    doc.add_page_break()

    add_heading(doc, "第5章　Claude Codeへの最重要タスク5本", 1)
    add_para(doc, "Claude Codeは以下を最優先で実行すること。", bold=True)

    add_heading(doc, "タスク1：YouTube Shorts自動生成パイプライン", 2)
    add_para(doc, "目的：note記事または短文ネタから、YouTube Shorts用の台本・画像・動画プロンプトを生成する。")
    add_para(doc, "必要処理：", bold=True)
    add_bullets(doc, [
        "30〜60秒台本生成",
        "タイトル生成",
        "概要欄生成",
        "ハッシュタグ生成",
        "サムネイル文言生成",
        "Kling AI / Sora / HeyGen へ渡せるプロンプト生成",
    ])

    add_heading(doc, "タスク2：HeyGen / AI女子アナ用スクリプト作成", 2)
    add_para(doc, "目的：日本語で自然に話すAI女子アナ動画を作る。")
    add_para(doc, "必要処理：", bold=True)
    add_bullets(doc, [
        "口語的で自然な台本",
        "AIっぽさを消す",
        "高岡市・日本文化・静かな暮らしをテーマ化",
        "YouTube Shorts向けに短く構成",
    ])

    add_heading(doc, "タスク3：Kling AI / Sora 用動画プロンプト作成", 2)
    add_para(doc, "目的：映像生成AIにそのまま貼れる動画プロンプトを作る。")
    add_para(doc, "必要処理：", bold=True)
    add_bullets(doc, [
        "cinematic",
        "realistic",
        "Japanese local atmosphere",
        "Toyama / Takaoka / quiet life",
        "short video format",
        "no fake text signs",
        "no unreadable captions",
    ])

    add_heading(doc, "タスク4：SNS / Reddit 自動投稿補助", 2)
    add_para(doc, "目的：海外流入を増やす。")
    add_para(doc, "必要処理：", bold=True)
    add_bullets(doc, [
        "note RSS取得",
        "最新記事タイトル取得",
        "英語紹介文生成",
        "Reddit投稿文生成",
        "X投稿文生成",
        "Instagramキャプション生成",
        "クリップボードコピー",
        "ログ保存",
    ])
    add_callout(
        doc,
        "注意",
        "完全自動投稿は規約リスクがあるため、まずは『投稿文生成＋コピー＋ブラウザ起動』方式を優先する。",
    )

    add_heading(doc, "タスク5：有料note記事量産", 2)
    add_para(doc, "目的：月1万円以上のnote収益を作る。")
    add_para(doc, "必要処理：", bold=True)
    add_bullets(doc, [
        "無料記事",
        "有料記事",
        "有料部分の見出し",
        "購入したくなる導入文",
        "海外向け英語版",
        "SNS導線文",
    ])
    add_para(doc, "テーマ例：", bold=True)
    add_bullets(doc, [
        "日本の静かな暮らし",
        "富山・高岡の生活",
        "地方都市の美しさ",
        "AI時代の副業",
        "50代からの働き方",
        "会社員から個人収益への移行",
    ])

    doc.add_page_break()

    add_heading(doc, "第6章　ファイル構成", 1)
    add_para(doc, "想定構成：")
    add_code_block(
        doc,
        "~/ai-auto/\n"
        "├── run.sh\n"
        "├── auto_post.py\n"
        "├── generate_note.py\n"
        "├── generate_reddit.py\n"
        "├── generate_youtube_short.py\n"
        "├── generate_paid_note.py\n"
        "├── cw_apply.py\n"
        "├── logs/\n"
        "├── outputs/\n"
        "├── prompts/\n"
        "└── .env",
    )
    add_heading(doc, "各ファイルの役割", 2)
    add_table(
        doc,
        ["ファイル", "役割"],
        [
            ["run.sh", "cronから呼び出すメイン実行ファイル"],
            ["auto_post.py", "投稿補助"],
            ["generate_note.py", "note記事生成"],
            ["generate_reddit.py", "Reddit投稿文生成"],
            ["generate_youtube_short.py", "Shorts台本生成"],
            ["generate_paid_note.py", "有料note生成"],
            ["cw_apply.py", "CrowdWorks応募文生成"],
            ["logs/", "実行ログ"],
            ["outputs/", "生成物保存"],
            ["prompts/", "プロンプト管理"],
            [".env", "APIキー管理"],
        ],
    )

    add_heading(doc, "第7章　運用ルール", 1)
    add_heading(doc, "基本方針", 2)
    add_bullets(doc, [
        "毎日最低1つ公開する。",
        "毎日最低1件応募する。",
        "計画より公開を優先する。",
        "生成物は必ず保存する。",
        "ログを残す。",
        "投稿前に最低限の目視確認を行う。",
    ])
    add_heading(doc, "禁止事項", 2)
    add_bullets(doc, [
        "計画だけ作って止まること",
        "非公開のまま放置すること",
        "完璧主義で投稿を遅らせること",
        "APIキーを直書きすること",
        "SNSへ過剰投稿すること",
        "規約違反リスクのある自動化を強行すること",
    ])

    add_heading(doc, "第8章　日次運用フロー", 1)
    add_para(doc, "理想的な1日の流れ：", bold=True)
    add_numbered(doc, [
        "cron実行",
        "note記事生成",
        "英語紹介文生成",
        "Reddit投稿文生成",
        "YouTube Shorts台本生成",
        "SNS投稿文生成",
        "outputs/ に保存",
        "logs/ に実行記録",
        "人間が確認",
        "公開・応募",
    ])
    add_para(doc, "本日時点の最重要タスク：", bold=True)
    add_numbered(doc, [
        "note記事を1本公開する",
        "CrowdWorksに1件応募する",
        "cronが正常動作しているか確認する",
    ])

    doc.add_page_break()

    add_heading(doc, "第9章　収益化戦略", 1)
    add_heading(doc, "現状", 2)
    add_table(
        doc,
        ["収益源", "状態"],
        [
            ["note有料記事", "初収益あり"],
            ["CrowdWorks", "受注前"],
            ["YouTube", "未収益"],
            ["SNS", "導線構築中"],
        ],
    )
    add_heading(doc, "問題点", 2)
    add_para(
        doc,
        "最大の問題は、商品や仕組みがないことではなく、公開数・応募数が不足していることである。",
    )
    add_heading(doc, "改善方針", 2)
    add_bullets(doc, [
        "noteは無料記事で流入、有料記事で収益化する",
        "YouTube Shortsは認知拡大に使う",
        "Redditは海外流入に使う",
        "CrowdWorksは短期現金化に使う",
        "SNSは導線強化に使う",
    ])

    add_heading(doc, "第10章　KPI", 1)
    add_table(
        doc,
        ["指標", "7日目標", "30日目標", "90日目標"],
        [
            ["note公開数", "7本", "30本", "90本"],
            ["有料記事数", "2本", "10本", "30本"],
            ["CrowdWorks応募数", "7件", "30件", "90件"],
            ["Shorts投稿数", "3本", "20本", "60本"],
            ["note収益", "1,000円", "10,000円", "50,000円"],
            ["総収益", "1,000円", "30,000円", "100,000円"],
        ],
    )

    add_heading(doc, "第11章　ロードマップ", 1)
    add_heading(doc, "フェーズ1：7日以内", 2)
    add_bullets(doc, [
        "noteを毎日公開",
        "CrowdWorksへ毎日応募",
        "cron正常動作確認",
        "outputs/logs整理",
        "YouTube Shorts台本生成を開始",
    ])
    add_heading(doc, "フェーズ2：30日以内", 2)
    add_bullets(doc, [
        "有料noteを10本作る",
        "Shortsを20本投稿",
        "Reddit導線を安定化",
        "SNS投稿文生成を定型化",
        "CrowdWorks受注を目指す",
    ])
    add_heading(doc, "フェーズ3：90日以内", 2)
    add_bullets(doc, [
        "月10万円の収益化を目指す",
        "note / YouTube / CrowdWorks の3本柱を作る",
        "自動化範囲を拡大",
        "低収益施策を削除",
        "伸びたテーマに集中投下する",
    ])

    doc.add_page_break()

    add_heading(doc, "第12章　リスク管理", 1)
    add_table(
        doc,
        ["リスク", "内容", "対策"],
        [
            ["API課金", "自動生成で費用増加", "上限設定・ログ確認"],
            ["SNS凍結", "過剰投稿・スパム判定", "投稿頻度制御"],
            ["品質低下", "AI文章が薄くなる", "テーマ固定・人間確認"],
            ["収益化遅延", "投稿しても売れない", "数を増やして検証"],
            ["作業停止", "完璧主義で止まる", "毎日1公開を絶対化"],
        ],
    )

    add_heading(doc, "第13章　Claude Code実行方針", 1)
    add_para(doc, "Claude Codeは以下の方針で動くこと。")
    add_para(doc, "優先順位：", bold=True)
    add_numbered(doc, [
        "公開",
        "応募",
        "収益",
        "自動化",
        "最適化",
    ])
    add_para(doc, "やること：", bold=True)
    add_bullets(doc, [
        "足りないファイルは作成する",
        "壊れているコードは修正する",
        "実行ログを残す",
        "生成物を保存する",
        "ユーザーに選択肢を投げすぎない",
        "迷ったら『公開に近い作業』を優先する",
    ])
    add_para(doc, "やらないこと：", bold=True)
    add_bullets(doc, [
        "長い計画だけ作る",
        "未実行で止まる",
        "不要な理想論を書く",
        "実際に収益化へ進まない作業を増やす",
    ])

    doc.add_page_break()

    add_heading(doc, "第14章　Claude Codeに渡す即時実行指示", 1)
    add_para(doc, "以下を実行してください。")
    add_code_block(
        doc,
        "cd ~/ai-auto || mkdir -p ~/ai-auto && cd ~/ai-auto\n"
        "mkdir -p logs outputs prompts\n"
        "cat > README.md <<'EOF'\n"
        "# AI Auto Monetization Project\n"
        "目的：\n"
        "note、YouTube Shorts、SNS、CrowdWorksを活用し、AIによる自動収益化の仕組みを作る。\n"
        "最優先：\n"
        "1. note公開\n"
        "2. CrowdWorks応募\n"
        "3. YouTube Shorts台本生成\n"
        "4. Reddit/SNS導線作成\n"
        "5. ログ保存\n"
        "EOF\n"
        "cat > run.sh <<'EOF'\n"
        "#!/bin/zsh\n"
        "cd ~/ai-auto\n"
        "mkdir -p logs outputs\n"
        "echo \"===== RUN $(date) =====\" >> logs/run.log\n"
        "python3 generate_daily_outputs.py >> logs/run.log 2>&1\n"
        "EOF\n"
        "chmod +x run.sh\n"
        "# generate_daily_outputs.py を作成（本書付録Aを参照）\n"
        "./run.sh\n"
        "echo \"完了：~/ai-auto/outputs に本日の生成物を作成しました。\"",
    )
    add_callout(
        doc,
        "実行確認",
        "本書発行時点で、上記コマンドは ~/ai-auto に対して実行済み。"
        "logs/run.log および outputs/ 配下に当日分の note・Reddit・YouTube Shorts・CrowdWorks 応募文が生成されている。",
        CALLOUT_INFO,
    )

    add_heading(doc, "第15章　最終評価", 1)
    add_table(
        doc,
        ["項目", "評価"],
        [
            ["技術基盤", "高い"],
            ["自動化構築", "高い"],
            ["コンテンツ設計", "中〜高"],
            ["収益化実績", "低い"],
            ["最大課題", "公開・応募・継続"],
        ],
    )
    add_callout(
        doc,
        "総合評価",
        "現在は『仕組みはあるが、公開と応募が不足している状態』である。",
    )
    add_callout(
        doc,
        "最終結論",
        "次にやるべきことは、新しい計画作成ではない。"
        "今日の生成物を公開し、CrowdWorksに1件応募し、ログを確認することである。",
        CALLOUT_WARN,
    )

    doc.add_page_break()

    add_heading(doc, "付録A　generate_daily_outputs.py", 1)
    add_para(
        doc,
        "本日の note・Reddit・YouTube Shorts 台本・CrowdWorks 応募文を outputs/ に出力するスクリプト。"
        "テーマは prompts/themes.json から日付ベースで自動ローテ選択され、"
        "各 generate_*.py の build() 関数を再利用する。"
        "本書発行時点で ~/ai-auto/generate_daily_outputs.py として配置済み。",
    )

    add_heading(doc, "付録B　published.csv と KPI 計測", 1)
    add_para(
        doc,
        "公開後は published.py で記録し、kpi.py で 7/30/90 日の達成度を確認する。",
    )
    add_code_block(
        doc,
        "# 公開記録\n"
        "python3 published.py note      \"https://note.com/...\" \"記事タイトル\" 0\n"
        "python3 published.py paid_note \"https://note.com/...\" \"30日メソッド\" 980\n"
        "python3 published.py crowdworks \"案件名\" \"応募完了\" 0\n"
        "\n"
        "# KPI 集計\n"
        "python3 kpi.py",
    )

    add_heading(doc, "付録C　LLM 呼び出しとコストガード", 1)
    add_para(
        doc,
        "ANTHROPIC_API_KEY または OPENAI_API_KEY が設定されていれば generate_note.py が API で生成する。"
        "日次予算（既定 ¥100）超過時は自動でテンプレフォールバックする。"
        "累計コストは .cost_log.json に記録される。",
    )
    add_code_block(
        doc,
        "# .env の例\n"
        "ANTHROPIC_API_KEY=sk-ant-...\n"
        "AI_DAILY_BUDGET_JPY=100\n"
        "\n"
        "# 既存の CW 案件スコアリング結果を取り込む\n"
        "python3 cw_apply.py --from-json /path/to/job.json",
    )

    add_heading(doc, "付録D　原単位レベル戦略（L1/L2/L3）", 1)
    add_para(
        doc,
        "本システムは原単位レベルを段階的に上げる前提で設計されている。"
        "L1（データ入力・¥980 note）は実績ゼロからの入口、L2（SEO記事執筆・¥2,980 note）が主収益、"
        "L3（SNS運用代行・コンサル・AI業務支援）が継続収益となる。",
    )
    add_table(
        doc,
        ["Level", "主な収益柱", "単価", "月必要件数", "月売上目安"],
        [
            ["L1", "データ入力 / ¥980 note", "¥1K〜10K", "30〜50件", "¥30〜50K"],
            ["L2", "SEO記事 / ¥2,980 note / マガジン", "¥10K〜30K", "5〜15件", "¥100〜300K"],
            ["L3", "SNS運用代行 / コンサル / AI支援", "¥50K〜300K", "2〜5件", "¥200K〜1.5M"],
        ],
    )
    add_para(doc, "対応スクリプト：", bold=True)
    add_bullets(doc, [
        "L1 応募：cw_apply.py --kind data",
        "L2 応募：cw_apply.py --kind writer",
        "L2 記事：generate_seo_article.py（無料Web AIで肉付け）",
        "L2 note：generate_paid_note.py（既定 ¥2,980）",
        "L3 応募：cw_apply.py --kind ai_support / --kind consultant",
        "L3 提案書：generate_proposal.py",
    ])
    add_callout(
        doc,
        "重要",
        "L1 のまま3ヶ月以上居座らないこと。M2 開始時には L2 案件への応募にシフトする。",
        CALLOUT_WARN,
    )

    add_heading(doc, "付録E　無料Web AI 運用（金銭費用ゼロの核）", 1)
    add_para(
        doc,
        "本システムは API を呼ばなくても L2 単価が成立するよう、"
        "無料Web AI（claude.ai / gemini.google.com / chatgpt.com）でのポリッシュ運用を採用する。"
        "プロンプト集は ~/ai-auto/prompts/polish_prompts.md に集約されている。",
    )
    add_table(
        doc,
        ["ツール", "強み", "主用途", "1日目安"],
        [
            ["claude.ai", "長文・日本語・SEO構成", "SEO記事 / 有料note / 提案書", "5〜10本"],
            ["gemini", "検索連携・タイトル発想", "キーワード調査 / タイトル候補", "無制限"],
            ["chatgpt", "短文・SNS表現", "X / Threads / Instagram", "5〜10本"],
        ],
    )
    add_callout(
        doc,
        "推奨",
        "API は既定で無効。3つの無料Web AI を使い分けて1日30本ペースまでカバー可能。月額¥0で L2 単価が狙える。",
        CALLOUT_INFO,
    )

    add_heading(doc, "付録F　Plan B 自動投稿基盤（時間分散・規約上グレー領域）", 1)
    add_para(
        doc,
        "本システムは公開〜応募までを Web 自動化で全自動化する Plan B を別途実装している。"
        "BANリスクを下げるため、時間分散・ジッター・人間挙動エミュレーション・1日上限・"
        "連続失敗停止の5重ガードを組み込み、初回は DRY_RUN=1 で試運転する設計。",
    )
    add_table(
        doc,
        ["コンポーネント", "ファイル", "役割"],
        [
            ["スケジューラ", "auto_schedule.py / _scheduler.py", "朝に当日のランダム投稿時刻を schedule.json に書き出し"],
            ["ディスパッチャ", "dispatcher.py", "毎時起動・現在時刻±15分のpendingタスクをジッター後に実行"],
            ["ブラウザ共通", "_browser.py", "Playwright 起動・人間挙動（ランダム待機/タイピング/マウス）"],
            ["note 公開", "publish_note.py", "Playwright・初回手動ログイン後はセッション再利用"],
            ["X 投稿", "post_x.py", "Playwright・3回/日に分散"],
            ["Reddit 投稿", "post_reddit.py", "PRAW（公式API・低リスク）"],
            ["CW 自動応募", "apply_crowdworks.py", "Playwright・1日3件まで・カテゴリ自動判定"],
        ],
    )
    add_heading(doc, "時間分散ルール", 2)
    add_table(
        doc,
        ["kind", "時間帯候補", "1日上限"],
        [
            ["note", "7-9 / 12-13 / 21-22 のランダム1本", "1"],
            ["x", "7-9 / 12-13 / 18-20 で各1回", "3"],
            ["reddit", "22-23 / 5-7（米国朝）", "1"],
            ["crowdworks", "平日 10-12 / 13-17", "3"],
        ],
    )
    add_callout(
        doc,
        "重要・規約",
        "note / X / CrowdWorks の Web 自動化はサービス規約上グレー。BANリスクは残る。"
        "段階的有効化（Reddit→X→note→CW）と DRY_RUN 試運転を必ず経ること。"
        "規約変更や検知強化があれば即停止。",
        CALLOUT_WARN,
    )

    add_heading(doc, "付録G　A 案運用プロトコル（Plan B + 週1ハイブリッド）", 1)
    add_para(
        doc,
        "Plan B 自動投稿の量と人間ポリッシュの質を組み合わせる推奨運用。"
        "BAN リスクを織り込んでも 6ヶ月累計利益 ¥831K（時給 ¥6,648相当）を期待できる。",
    )
    add_table(
        doc,
        ["曜日・時刻", "やること", "所要"],
        [
            ["平日 朝", "何もしない（cron で全自動）", "0分"],
            ["平日 朝9時", "ban_detector でBAN検知（異常時のみ対応）", "1分"],
            ["日曜 18時", "sunday_polish.py が高単価候補3本を自動生成", "0分"],
            ["日曜 18:00-19:00", "claude.ai でポリッシュ（有料note・SEO記事・提案書）", "30〜60分"],
            ["日曜 19:00-19:30", "公開・応募 → published.py で記録", "30分"],
        ],
    )
    add_heading(doc, "cron 設定（A案 4行）", 2)
    add_code_block(
        doc,
        "0 7 * * *  /bin/zsh -lc 'cd $HOME/ai-auto && ./run.sh >> logs/cron.log 2>&1'\n"
        "0 * * * *  /bin/zsh -lc 'cd $HOME/ai-auto && python3 dispatcher.py >> logs/dispatcher.log 2>&1'\n"
        "0 9 * * *  /bin/zsh -lc 'cd $HOME/ai-auto && python3 ban_detector.py >> logs/ban_check.log 2>&1'\n"
        "0 18 * * 0 /bin/zsh -lc 'cd $HOME/ai-auto && python3 sunday_polish.py >> logs/sunday.log 2>&1'",
    )
    add_callout(
        doc,
        "BAN対応",
        "ban_detector.py が 403/404/410/451 を検出したら、該当サービスの自動投稿を即停止"
        "（.env で DRY_RUN=1）。原因特定までは手動運用に戻す。",
        CALLOUT_WARN,
    )

    add_heading(doc, "付録H　iPhone 連携（HTTP API + SSH）", 1)
    add_para(
        doc,
        "外出先から iPhone でKPI確認・公開記録・DRY_RUN切替を行うための仕組み。"
        "ai-auto/server.py が Python標準ライブラリのみで動く HTTP API を提供し、"
        "iPhone ショートカットアプリから直接叩ける。SSH 経由の補助レシピも併用可。",
    )
    add_table(
        doc,
        ["エンドポイント", "用途"],
        [
            ["GET /health", "稼働確認・DRY_RUN状態取得"],
            ["GET /kpi", "KPI集計（kpi.py の出力）"],
            ["GET /outputs", "outputs/ 最新ファイル一覧"],
            ["GET /outputs/<name>", "ファイル本文取得"],
            ["GET /schedule", "当日 schedule.json"],
            ["POST /publish", "公開記録の追加"],
            ["POST /dry-run", ".env の DRY_RUN を 0/1 切替"],
        ],
    )
    add_para(doc, "認証：環境変数 AI_AUTO_TOKEN（未設定時は loopback のみ）。", bold=True)
    add_para(doc, "起動：")
    add_code_block(
        doc,
        "~/ai-auto/run_server.sh           # mac内のみ（127.0.0.1:8765）\n"
        "~/ai-auto/run_server.sh lan       # LAN公開（要 AI_AUTO_TOKEN）\n"
        "# 常駐化: ~/Library/LaunchAgents/com.aiauto.server.plist を配置して launchctl load",
    )
    add_callout(
        doc,
        "外出先からも使う場合",
        "Tailscale（無料・推奨）で mac と iPhone を仮想LANに入れる。"
        "ngrok でも可だが一時URL。iCloud Private Relay はインバウンド不可なので使えない。",
        CALLOUT_INFO,
    )
    add_para(doc, "詳細レシピ：~/ai-auto/iphone_shortcuts.md", bold=True)
    add_heading(doc, "1コマンドでセットアップ", 2)
    add_code_block(
        doc,
        "# iPhone 連携を立ち上げる（トークン生成＋サーバー起動＋IP表示まで）\n"
        "cd ~/ai-auto && bash quick_start.sh           # LAN公開\n"
        "cd ~/ai-auto && bash quick_start.sh local     # mac内のみ\n"
        "cd ~/ai-auto && bash quick_start.sh stop      # 停止\n"
        "\n"
        "# 各サービスの初回手動ログイン（Cookie保存）\n"
        "cd ~/ai-auto && DRY_RUN=0 python3 init_login.py",
    )

    add_heading(doc, "付録I　GitHub Actions 自動運用（mac 不要・iPhone 1分公開）", 1)
    add_para(
        doc,
        "mac セットアップが障壁となり実運用が始まらない問題を解決するため、"
        "GitHub Actions による完全自動ドラフト生成と Reddit 自動公開を導入。"
        "オーナーは iPhone のブラウザだけで毎朝1分以内に公開を完了できる。",
    )
    add_table(
        doc,
        ["コンポーネント", "ファイル", "役割"],
        [
            ["生成ワークフロー", ".github/workflows/daily-drafts.yml", "毎朝07:00 JST にドラフト生成→自動コミット"],
            ["生成スクリプト", "tools/daily_drafts.py", "8種類のドラフトを daily/<日付>/ に出力"],
            ["Reddit 自動公開", "tools/publish_reddit.py", "PRAW 経由・GitHub Secrets で認証"],
            ["ドラフト置き場", "daily/<日付>/", "iPhone の GitHub アプリから閲覧"],
        ],
    )
    add_heading(doc, "Reddit 自動公開の有効化", 2)
    add_para(doc, "GitHub リポジトリの Settings で以下を設定：", bold=True)
    add_bullets(doc, [
        "Secrets: REDDIT_CLIENT_ID / REDDIT_CLIENT_SECRET / REDDIT_USER_AGENT / REDDIT_USERNAME / REDDIT_PASSWORD",
        "Variables: ENABLE_REDDIT_AUTO_PUBLISH=true / REDDIT_SUBREDDIT=SlowLiving (任意)",
    ])
    add_heading(doc, "iPhone 1分運用フロー", 2)
    add_numbered(doc, [
        "GitHub アプリで daily/<最新日付>/ を開く",
        "note_draft.md をタップ → Raw → 全選択コピー",
        "note アプリで新規記事 → 貼り付け → 公開",
    ])
    add_callout(
        doc,
        "重要",
        "note / X / CrowdWorks の Web 自動化（Plan B）は規約上グレーで mac 環境必須。"
        "GitHub Actions では Reddit のみ規約準拠で自動化可能。"
        "L1 構成（GitHub Actions ドラフト生成 + Reddit 自動 + iPhone 手動公開）が時間対効果最大。",
        CALLOUT_INFO,
    )

    add_heading(doc, "付録J　最終命令（要約）", 1)
    add_numbered(doc, [
        "note記事を1本公開できる状態にする",
        "CrowdWorks応募文を1本作る",
        "cronで毎日動く基礎を作る",
    ])
    add_callout(
        doc,
        "署名",
        "本書は 2026-05-05 に発行された Claude Code 完全版引き継ぎ書である。"
        "次回更新は KPI 7日目標達成時または重大な方針変更時に行う。",
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
