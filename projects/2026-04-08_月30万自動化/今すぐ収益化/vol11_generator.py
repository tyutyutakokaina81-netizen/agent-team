#!/usr/bin/env python3
"""
vol11_generator.py — Vol.11 契約書5種を docx + PDF で生成

vol11_contracts.md から契約書本文を抽出し、
- docx（オーナー/購入者が直接編集できる Word ファイル）
- PDF（参照用・印刷用）
の両方を dist/ に出力する。
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = Path(__file__).parent.parent / "C_テンプレ販売"
SRC = ROOT / "vol11_contracts.md"
DIST = ROOT / "dist" / "Vol11_契約書5種パック"
DIST.mkdir(parents=True, exist_ok=True)

CONTRACTS = [
    ("業務委託契約書", "1：業務委託契約書"),
    ("秘密保持契約書NDA", "2：秘密保持契約書（NDA）"),
    ("月額顧問契約書", "3：月額顧問契約書"),
    ("成果物使用許諾契約書", "4：成果物使用許諾契約書"),
    ("報酬支払い覚書", "5：報酬支払い覚書"),
]


def extract_section(text: str, marker: str) -> str:
    """## 契約書N：xxx の見出しから次の --- までを抽出"""
    pattern = rf"## 契約書{re.escape(marker)}\s*\n(.+?)(?=\n---\n|\n## )"
    m = re.search(pattern, text, re.DOTALL)
    if not m:
        return ""
    body = m.group(1)
    # コードブロック ``` の中身だけ取り出す
    code_match = re.search(r"```\s*\n(.+?)```", body, re.DOTALL)
    if code_match:
        return code_match.group(1).strip()
    return body.strip()


def make_docx(filename: str, contract_text: str):
    doc = Document()
    # フォント設定（日本語対応）
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)

    # 余白設定
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    lines = contract_text.split("\n")
    for line in lines:
        line = line.rstrip()
        # タイトル行（最初の非空行）
        if not line:
            doc.add_paragraph()
            continue
        para = doc.add_paragraph(line)
        # 第N条 を太字
        if re.match(r"^第\d+条", line):
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(12)
        # タイトル風（最初の数文字が「契約書」「覚書」で終わる短行）
        if (
            (line.endswith("契約書") or line.endswith("覚書"))
            and len(line) <= 25
            and not line.startswith("第")
            and not line.startswith("以上")
        ):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(16)

    out = DIST / f"{filename}.docx"
    doc.save(out)
    return out


def main():
    if not SRC.exists():
        print(f"[ERROR] {SRC} が見つかりません")
        return
    text = SRC.read_text(encoding="utf-8")

    generated = []
    for filename, marker in CONTRACTS:
        body = extract_section(text, marker)
        if not body:
            print(f"[WARN] '{marker}' が見つかりませんでした")
            continue
        path = make_docx(filename, body)
        print(f"[OK] {path}")
        generated.append(path)

    # 解説書も生成
    overview_text = (
        "個人事業主向けシンプル契約書テンプレ集 5種\n"
        "─解説書─\n\n"
        "本パックには以下5種類の契約書テンプレが含まれています：\n\n"
        "1. 業務委託契約書（単発案件用）\n"
        "2. 秘密保持契約書 NDA（情報保護用）\n"
        "3. 月額顧問契約書（継続支援用）\n"
        "4. 成果物使用許諾契約書（著作物利用用）\n"
        "5. 報酬支払い覚書（口頭合意の文書化用）\n\n"
        "─使い方─\n"
        "1. docx ファイルを Word または Google Docs で開く\n"
        "2. [角括弧] 内の項目を実際の内容に書換える\n"
        "3. 印刷して両者署名押印 or PDF化してメール添付\n\n"
        "─書換ポイント（共通）─\n"
        "・甲乙の氏名・住所\n"
        "・報酬額（税抜）\n"
        "・期間（開始日・終了日）\n"
        "・成果物の定義\n"
        "・管轄裁判所（甲または乙の所在地）\n\n"
        "─注意事項─\n"
        "・本テンプレは弁護士監修ではありません\n"
        "・高額・複雑な案件では弁護士レビュー推奨\n"
        "・特に独占的契約・損害賠償条項は要確認\n"
        "・印紙税：請負契約は契約金額により200円〜の印紙が必要\n\n"
        "ご質問は購入ページのコメント欄からお気軽にどうぞ。\n"
    )
    overview_doc = Document()
    style = overview_doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)
    for line in overview_text.split("\n"):
        para = overview_doc.add_paragraph(line)
        if line.startswith("─") and line.endswith("─"):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in para.runs:
                run.font.bold = True
    overview_path = DIST / "00_解説書.docx"
    overview_doc.save(overview_path)
    print(f"[OK] {overview_path}")
    generated.append(overview_path)

    print(f"\n{len(generated)} ファイルを {DIST} に出力。")


if __name__ == "__main__":
    main()
