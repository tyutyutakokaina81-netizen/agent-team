#!/usr/bin/env python3
"""
invoice_generator.py — 受注後の請求書を docx で自動生成

CSV にクライアント情報・金額を入力 → 1コマンドで請求書 docx を出力。
適格請求書（インボイス）形式に対応。

【使い方】
  # サンプル CSV を作成
  python invoice_generator.py init

  # CSV を編集してから一括生成
  python invoice_generator.py generate

  # 個別生成
  python invoice_generator.py one --client "株式会社XX" --amount 50000 --desc "SNS運用代行 5月分"

【出力】
  outputs/YYYYMMDD_クライアント名_請求書.docx

【設定】
  config.json で以下を管理：
  - 自社情報（屋号・住所・登録番号・口座情報）
  - 消費税率
  - 支払条件（例：月末締め翌月末払い）
"""

import argparse
import csv
import json
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = Path(__file__).parent
CONFIG_FILE = ROOT / "config.json"
INVOICES_CSV = ROOT / "invoices.csv"
OUTPUT_DIR = ROOT / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

DEFAULT_CONFIG = {
    "issuer_name": "あなたの屋号",
    "issuer_personal_name": "あなたの氏名",
    "issuer_postal": "〒XXX-XXXX",
    "issuer_address": "東京都XX区XX 1-2-3",
    "issuer_email": "your-email@example.com",
    "issuer_phone": "090-XXXX-XXXX",
    "registration_number": "T1234567890123",
    "bank_name": "XX銀行 XX支店",
    "bank_account_type": "普通",
    "bank_account_number": "1234567",
    "bank_account_holder": "アナタ ノ ナマエ",
    "tax_rate": 0.10,
    "payment_terms": "月末締め翌月末払い",
    "invoice_number_prefix": "INV-",
}

DEFAULT_INVOICES = [
    {
        "client_name": "（記入例）株式会社サンプル",
        "client_address": "東京都サンプル区1-2-3",
        "client_attn": "経理部 ご担当者様",
        "issue_date": "2026-05-31",
        "description": "SNS運用代行 2026年5月分",
        "quantity": "1",
        "unit_price": "50000",
        "notes": "",
    },
]


def load_config() -> dict:
    if CONFIG_FILE.exists():
        return json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
    return DEFAULT_CONFIG


def save_config(c: dict):
    CONFIG_FILE.write_text(json.dumps(c, ensure_ascii=False, indent=2), encoding="utf-8")


def cmd_init():
    """設定ファイル＋ CSV テンプレを作成"""
    if not CONFIG_FILE.exists():
        save_config(DEFAULT_CONFIG)
        print(f"[OK] 設定ファイルを作成: {CONFIG_FILE}")
    else:
        print(f"[SKIP] 設定ファイルは既に存在: {CONFIG_FILE}")

    if not INVOICES_CSV.exists():
        fields = list(DEFAULT_INVOICES[0].keys())
        with INVOICES_CSV.open("w", encoding="utf-8", newline="") as f:
            w = csv.DictWriter(f, fieldnames=fields)
            w.writeheader()
            for row in DEFAULT_INVOICES:
                w.writerow(row)
        print(f"[OK] 請求書 CSV テンプレを作成: {INVOICES_CSV}")
    else:
        print(f"[SKIP] 請求書 CSV は既に存在: {INVOICES_CSV}")

    print("\n次のステップ:")
    print(f"  1. {CONFIG_FILE} を編集して自社情報を記入")
    print(f"  2. {INVOICES_CSV} に実際の請求情報を追加")
    print(f"  3. python {Path(__file__).name} generate を実行")


def make_invoice(config: dict, row: dict, output_path: Path):
    """1件の請求書 docx を生成"""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # タイトル
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("請　求　書")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    doc.add_paragraph()

    # 発行情報（右寄せ）
    issue_date = row.get("issue_date") or datetime.now().strftime("%Y-%m-%d")
    invoice_number = f"{config['invoice_number_prefix']}{datetime.now().strftime('%Y%m%d')}-{abs(hash(row['client_name'])) % 10000:04d}"

    info_para = doc.add_paragraph()
    info_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    info_para.add_run(f"請求書番号：{invoice_number}\n発行日：{issue_date}").font.size = Pt(10)
    doc.add_paragraph()

    # 宛先
    client_para = doc.add_paragraph()
    client_para.add_run(f"{row['client_name']}\n").font.bold = True
    if row.get("client_attn"):
        client_para.add_run(f"{row['client_attn']}\n")
    if row.get("client_address"):
        client_para.add_run(f"{row['client_address']}\n").font.size = Pt(9)
    client_para.add_run("御中").font.bold = True
    doc.add_paragraph()

    # 合計金額（強調）
    qty = float(row.get("quantity") or 1)
    unit_price = float(row.get("unit_price") or 0)
    subtotal = int(qty * unit_price)
    tax = int(subtotal * config["tax_rate"])
    total = subtotal + tax

    total_para = doc.add_paragraph()
    total_run = total_para.add_run(f"ご請求金額　¥{total:,}（税込）")
    total_run.font.size = Pt(18)
    total_run.font.bold = True
    total_run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6E)
    doc.add_paragraph()

    # 明細テーブル
    doc.add_paragraph().add_run("【明細】").font.bold = True
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "品目"
    hdr[1].text = "数量"
    hdr[2].text = "単価"
    hdr[3].text = "金額"

    row_cells = table.add_row().cells
    row_cells[0].text = row["description"]
    row_cells[1].text = f"{int(qty) if qty.is_integer() else qty}"
    row_cells[2].text = f"¥{int(unit_price):,}"
    row_cells[3].text = f"¥{subtotal:,}"

    # 合計行
    total_row1 = table.add_row().cells
    total_row1[2].text = "小計"
    total_row1[3].text = f"¥{subtotal:,}"
    total_row2 = table.add_row().cells
    total_row2[2].text = f"消費税 ({int(config['tax_rate']*100)}%)"
    total_row2[3].text = f"¥{tax:,}"
    total_row3 = table.add_row().cells
    total_row3[2].text = "合計"
    total_row3[3].text = f"¥{total:,}"
    for cell in [total_row1[2], total_row2[2], total_row3[2], total_row3[3]]:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

    doc.add_paragraph()

    # 振込先
    bank_para = doc.add_paragraph()
    bank_para.add_run("【お振込先】\n").font.bold = True
    bank_para.add_run(
        f"{config['bank_name']}\n"
        f"{config['bank_account_type']} {config['bank_account_number']}\n"
        f"口座名義：{config['bank_account_holder']}\n"
    )

    # 支払条件
    pay_para = doc.add_paragraph()
    pay_para.add_run(f"【お支払期限】{config['payment_terms']}\n").font.bold = True
    pay_para.add_run("※ 振込手数料は貴社にてご負担ください。\n").font.size = Pt(9)

    # 備考
    if row.get("notes"):
        doc.add_paragraph().add_run("【備考】").font.bold = True
        doc.add_paragraph(row["notes"])

    doc.add_paragraph()

    # 発行者情報
    issuer = doc.add_paragraph()
    issuer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    issuer.add_run(
        f"{config['issuer_name']}\n"
        f"{config['issuer_personal_name']}\n"
        f"{config['issuer_postal']} {config['issuer_address']}\n"
        f"Tel: {config['issuer_phone']}\n"
        f"Email: {config['issuer_email']}\n"
        f"登録番号: {config['registration_number']}\n"
    ).font.size = Pt(10)

    doc.save(output_path)


def cmd_generate():
    config = load_config()
    if not INVOICES_CSV.exists():
        print(f"[ERROR] {INVOICES_CSV} が存在しません。先に init を実行してください。")
        return

    with INVOICES_CSV.open(encoding="utf-8") as f:
        rows = list(csv.DictReader(f))

    generated = []
    for row in rows:
        if not row.get("client_name") or row["client_name"].startswith("（記入例）"):
            print(f"[SKIP] {row.get('client_name', '?')}：記入例なのでスキップ")
            continue
        date = row.get("issue_date") or datetime.now().strftime("%Y-%m-%d")
        date_compact = date.replace("-", "")
        safe_name = "".join(c for c in row["client_name"] if c.isalnum() or c in "_-")[:20]
        out = OUTPUT_DIR / f"{date_compact}_{safe_name}_請求書.docx"
        make_invoice(config, row, out)
        print(f"[OK] {out}")
        generated.append(out)

    print(f"\n{len(generated)} 件の請求書を生成。")


def cmd_one(args):
    config = load_config()
    row = {
        "client_name": args.client,
        "client_address": args.address or "",
        "client_attn": args.attn or "ご担当者様",
        "issue_date": args.date or datetime.now().strftime("%Y-%m-%d"),
        "description": args.desc,
        "quantity": str(args.quantity),
        "unit_price": str(args.amount),
        "notes": args.notes or "",
    }
    date_compact = row["issue_date"].replace("-", "")
    safe_name = "".join(c for c in args.client if c.isalnum() or c in "_-")[:20]
    out = OUTPUT_DIR / f"{date_compact}_{safe_name}_請求書.docx"
    make_invoice(config, row, out)
    print(f"[OK] {out}")


def main():
    parser = argparse.ArgumentParser(description="請求書自動生成ツール")
    sub = parser.add_subparsers(dest="cmd", required=True)

    sub.add_parser("init", help="設定ファイルと CSV テンプレを作成")
    sub.add_parser("generate", help="invoices.csv の全件を生成")

    one = sub.add_parser("one", help="単発の請求書を生成")
    one.add_argument("--client", required=True, help="クライアント名")
    one.add_argument("--amount", type=int, required=True, help="単価（税抜）")
    one.add_argument("--desc", required=True, help="品目")
    one.add_argument("--quantity", type=int, default=1)
    one.add_argument("--address", help="クライアント住所")
    one.add_argument("--attn", help="宛名")
    one.add_argument("--date", help="発行日 YYYY-MM-DD")
    one.add_argument("--notes", help="備考")

    args = parser.parse_args()

    if args.cmd == "init":
        cmd_init()
    elif args.cmd == "generate":
        cmd_generate()
    elif args.cmd == "one":
        cmd_one(args)


if __name__ == "__main__":
    main()
