#!/usr/bin/env python3
"""
quote_receipt_generator.py — 見積書・領収書を docx で自動生成

invoice_generator.py と同じ config.json を共有する補助ツール。

【使い方】
  # 見積書（受注前）
  python quote_receipt_generator.py quote \
    --client "株式会社XX" --amount 150000 --desc "SEO記事10本"

  # 領収書（入金確認後）
  python quote_receipt_generator.py receipt \
    --client "株式会社XX" --amount 150000 --desc "SEO記事10本" --paid-date 2026-05-31

【出力】
  outputs/YYYYMMDD_クライアント名_見積書.docx
  outputs/YYYYMMDD_クライアント名_領収書.docx
"""

import argparse
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = Path(__file__).parent
CONFIG_FILE = ROOT / "config.json"
OUTPUT_DIR = ROOT / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)


def load_config() -> dict:
    if not CONFIG_FILE.exists():
        print(f"[ERROR] {CONFIG_FILE} が無い。先に invoice_generator.py init を実行してください。")
        raise SystemExit(1)
    return json.loads(CONFIG_FILE.read_text(encoding="utf-8"))


def add_title(doc, label: str):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(label)
    run.font.size = Pt(22)
    run.font.bold = True


def add_client_block(doc, client: str, address: str = "", attn: str = ""):
    para = doc.add_paragraph()
    para.add_run(f"{client}\n").font.bold = True
    if attn:
        para.add_run(f"{attn}\n")
    if address:
        para.add_run(f"{address}\n").font.size = Pt(9)
    para.add_run("御中").font.bold = True


def add_total_block(doc, total: int, label: str = "ご請求金額"):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}　¥{total:,}（税込）")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6E)


def add_detail_table(doc, desc: str, qty: int, unit_price: int, tax_rate: float):
    subtotal = qty * unit_price
    tax = int(subtotal * tax_rate)
    total = subtotal + tax

    doc.add_paragraph().add_run("【明細】").font.bold = True
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "品目"
    hdr[1].text = "数量"
    hdr[2].text = "単価"
    hdr[3].text = "金額"

    row = table.add_row().cells
    row[0].text = desc
    row[1].text = str(qty)
    row[2].text = f"¥{unit_price:,}"
    row[3].text = f"¥{subtotal:,}"

    table.add_row().cells[2].text = "小計"; table.rows[-1].cells[3].text = f"¥{subtotal:,}"
    table.add_row().cells[2].text = f"消費税 ({int(tax_rate*100)}%)"; table.rows[-1].cells[3].text = f"¥{tax:,}"
    table.add_row().cells[2].text = "合計"; table.rows[-1].cells[3].text = f"¥{total:,}"

    return subtotal, tax, total


def add_issuer_block(doc, config: dict, with_seal: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    seal = "（印）" if with_seal else ""
    p.add_run(
        f"{config['issuer_name']} {seal}\n"
        f"{config['issuer_personal_name']}\n"
        f"{config['issuer_postal']} {config['issuer_address']}\n"
        f"Tel: {config['issuer_phone']}\n"
        f"Email: {config['issuer_email']}\n"
        f"登録番号: {config['registration_number']}\n"
    ).font.size = Pt(10)


def setup_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    return doc


# ─────────────────────────────────
# 見積書
# ─────────────────────────────────

def make_quote(args, config: dict):
    doc = setup_doc()
    add_title(doc, "御　見　積　書")
    doc.add_paragraph()

    issue_date = args.date or datetime.now().strftime("%Y-%m-%d")
    quote_number = f"Q-{datetime.now().strftime('%Y%m%d')}-{abs(hash(args.client)) % 10000:04d}"
    valid_until = args.valid_until or "発行日から30日"

    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    info.add_run(
        f"見積番号：{quote_number}\n"
        f"発行日：{issue_date}\n"
        f"見積有効期限：{valid_until}\n"
    ).font.size = Pt(10)
    doc.add_paragraph()

    add_client_block(doc, args.client, args.address or "", args.attn or "ご担当者様")
    doc.add_paragraph()
    doc.add_paragraph("下記のとおりお見積もり申し上げます。")
    doc.add_paragraph()

    subtotal, tax, total = add_detail_table(doc, args.desc, args.quantity, args.amount, config["tax_rate"])
    doc.add_paragraph()
    add_total_block(doc, total, label="お見積金額")
    doc.add_paragraph()

    if args.notes:
        doc.add_paragraph().add_run("【備考】").font.bold = True
        doc.add_paragraph(args.notes)

    doc.add_paragraph().add_run(
        "※ 上記金額は消費税込みです。\n"
        f"※ 本見積書の有効期限は {valid_until} です。\n"
        "※ 受注後、別途請求書を発行いたします。\n"
    ).font.size = Pt(9)

    add_issuer_block(doc, config, with_seal=True)

    safe_name = "".join(c for c in args.client if c.isalnum() or c in "_-")[:20]
    out = OUTPUT_DIR / f"{issue_date.replace('-','')}_{safe_name}_見積書.docx"
    doc.save(out)
    print(f"[OK] {out}")


# ─────────────────────────────────
# 領収書
# ─────────────────────────────────

def make_receipt(args, config: dict):
    doc = setup_doc()
    add_title(doc, "領　収　書")
    doc.add_paragraph()

    paid_date = args.paid_date or datetime.now().strftime("%Y-%m-%d")
    receipt_number = f"R-{datetime.now().strftime('%Y%m%d')}-{abs(hash(args.client)) % 10000:04d}"

    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    info.add_run(
        f"領収番号：{receipt_number}\n"
        f"発行日：{paid_date}\n"
    ).font.size = Pt(10)
    doc.add_paragraph()

    add_client_block(doc, args.client, args.address or "", args.attn or "ご担当者様")
    doc.add_paragraph()

    subtotal = args.amount * args.quantity
    tax = int(subtotal * config["tax_rate"])
    total = subtotal + tax
    add_total_block(doc, total, label="領収金額")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run(
        f"上記金額を {args.method or '銀行振込'} にて\n"
        f"{paid_date} に正に領収いたしました。\n"
    )
    doc.add_paragraph()

    doc.add_paragraph().add_run("【内訳】").font.bold = True
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid"
    table.rows[0].cells[0].text = "品目"
    table.rows[0].cells[1].text = "金額"
    row = table.add_row().cells
    row[0].text = args.desc
    row[1].text = f"¥{subtotal:,}"
    row = table.add_row().cells
    row[0].text = f"消費税 ({int(config['tax_rate']*100)}%)"
    row[1].text = f"¥{tax:,}"
    row = table.add_row().cells
    row[0].text = "合計"
    row[1].text = f"¥{total:,}"

    doc.add_paragraph()
    doc.add_paragraph().add_run(
        "※ 本領収書は適格請求書の代替として税務処理にご利用いただけます。\n"
        "※ 印紙税：5万円以上の現金受領は収入印紙が必要（振込の場合は不要）。\n"
    ).font.size = Pt(9)

    add_issuer_block(doc, config, with_seal=True)

    safe_name = "".join(c for c in args.client if c.isalnum() or c in "_-")[:20]
    out = OUTPUT_DIR / f"{paid_date.replace('-','')}_{safe_name}_領収書.docx"
    doc.save(out)
    print(f"[OK] {out}")


# ─────────────────────────────────
# CLI
# ─────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="見積書・領収書 自動生成")
    sub = parser.add_subparsers(dest="cmd", required=True)

    q = sub.add_parser("quote", help="見積書を生成")
    q.add_argument("--client", required=True)
    q.add_argument("--amount", type=int, required=True, help="単価（税抜）")
    q.add_argument("--desc", required=True)
    q.add_argument("--quantity", type=int, default=1)
    q.add_argument("--address")
    q.add_argument("--attn")
    q.add_argument("--date", help="発行日 YYYY-MM-DD")
    q.add_argument("--valid-until", help="見積有効期限")
    q.add_argument("--notes")

    r = sub.add_parser("receipt", help="領収書を生成")
    r.add_argument("--client", required=True)
    r.add_argument("--amount", type=int, required=True)
    r.add_argument("--desc", required=True)
    r.add_argument("--quantity", type=int, default=1)
    r.add_argument("--address")
    r.add_argument("--attn")
    r.add_argument("--paid-date", help="入金日 YYYY-MM-DD")
    r.add_argument("--method", help="支払方法（例：銀行振込）")

    args = parser.parse_args()
    config = load_config()
    if args.cmd == "quote":
        make_quote(args, config)
    elif args.cmd == "receipt":
        make_receipt(args, config)


if __name__ == "__main__":
    main()
